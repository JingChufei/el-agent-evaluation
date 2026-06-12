from __future__ import annotations

import csv
import mimetypes
import re
import zipfile
from pathlib import Path
from typing import Any

from openpyxl import load_workbook

from .common import is_blankish, sha256_file, split_multiline
from .paths import as_posix_relative


def _mime_type(path: Path) -> str:
    guessed, _ = mimetypes.guess_type(path.name)
    return guessed or "application/octet-stream"


def _read_csv_preflight(path: Path) -> dict[str, Any]:
    errors: list[str] = []
    encoding = "utf-8-sig"
    text = ""
    for candidate in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            text = path.read_text(encoding=candidate)
            encoding = candidate
            break
        except UnicodeDecodeError as exc:
            errors.append(f"{candidate}: {exc}")
    rows = []
    if text:
        reader = csv.reader(text.splitlines())
        rows = list(reader)
    return {
        "kind": "csv",
        "readable": bool(text or path.stat().st_size == 0),
        "encoding": encoding,
        "rows": len(rows),
        "max_columns": max((len(row) for row in rows), default=0),
        "sample_header": rows[0] if rows else [],
        "errors": errors[-1:] if not text and path.stat().st_size else [],
    }


def _read_xlsx_preflight(path: Path) -> dict[str, Any]:
    try:
        workbook = load_workbook(path, read_only=True, data_only=True)
        sheets = [
            {"name": sheet.title, "max_row": sheet.max_row, "max_column": sheet.max_column}
            for sheet in workbook.worksheets
        ]
        workbook.close()
        return {"kind": "xlsx", "readable": True, "sheets": sheets}
    except Exception as exc:  # pragma: no cover - exact parser errors vary by file
        return {"kind": "xlsx", "readable": False, "error": str(exc)}


def _read_pdf_preflight(path: Path) -> dict[str, Any]:
    try:
        from pypdf import PdfReader  # type: ignore

        reader = PdfReader(str(path))
        return {"kind": "pdf", "readable": True, "pages": len(reader.pages)}
    except Exception as exc:
        header_ok = path.read_bytes()[:4] == b"%PDF"
        return {"kind": "pdf", "readable": header_ok, "parser": "header_only", "error": str(exc)}


def _read_docx_preflight(path: Path) -> dict[str, Any]:
    try:
        from docx import Document  # type: ignore

        document = Document(str(path))
        return {"kind": "docx", "readable": True, "paragraphs": len(document.paragraphs), "tables": len(document.tables)}
    except Exception as exc:
        try:
            with zipfile.ZipFile(path) as archive:
                archive.testzip()
                names = archive.namelist()
            return {"kind": "docx", "readable": True, "parser": "zip_only", "entries": len(names)}
        except Exception as zip_exc:
            return {"kind": "docx", "readable": False, "error": f"{exc}; zip={zip_exc}"}


def _read_pptx_preflight(path: Path) -> dict[str, Any]:
    try:
        with zipfile.ZipFile(path) as archive:
            archive.testzip()
            slide_names = [
                name
                for name in archive.namelist()
                if re.fullmatch(r"ppt/slides/slide\d+\.xml", name)
            ]
        return {"kind": "pptx", "readable": True, "slides": len(slide_names)}
    except Exception as exc:
        return {"kind": "pptx", "readable": False, "error": str(exc)}


def preflight_file(path: Path) -> dict[str, Any]:
    suffix = path.suffix.lower()
    if suffix == ".csv":
        return _read_csv_preflight(path)
    if suffix == ".xlsx":
        return _read_xlsx_preflight(path)
    if suffix == ".pdf":
        return _read_pdf_preflight(path)
    if suffix == ".docx":
        return _read_docx_preflight(path)
    if suffix == ".pptx":
        return _read_pptx_preflight(path)
    return {"kind": suffix.lstrip(".") or "unknown", "readable": path.is_file()}


def build_attachment_manifest(attachments_root: Path, *, repo_root: Path | None = None) -> list[dict[str, Any]]:
    entries: list[dict[str, Any]] = []
    if not attachments_root.exists():
        return entries
    repo_root = repo_root or attachments_root.parent
    for path in sorted(attachments_root.rglob("*")):
        if not path.is_file():
            continue
        if path.name.startswith("~$") or path.name == ".DS_Store":
            continue
        stat = path.stat()
        entries.append(
            {
                "basename": path.name,
                "relative_path": str(path.relative_to(attachments_root)),
                "repo_relative_path": as_posix_relative(path, repo_root),
                "absolute_path": str(path.resolve()),
                "extension": path.suffix.lower(),
                "size_bytes": stat.st_size,
                "mime_type": _mime_type(path),
                "sha256": sha256_file(path),
                "preflight": preflight_file(path),
            }
        )
    return entries


def _entry_to_attachment(entry: dict[str, Any]) -> dict[str, Any]:
    return {
        "name": entry["basename"],
        "relative_path": entry["relative_path"],
        "original_relative_path": entry.get("repo_relative_path", entry["relative_path"]),
        "original_path": entry["absolute_path"],
        "mime_type": entry["mime_type"],
        "extension": entry["extension"],
        "size_bytes": entry["size_bytes"],
        "sha256": entry["sha256"],
        "preflight": entry["preflight"],
    }


def resolve_attachment_cell(cell_value: str, manifest: list[dict[str, Any]]) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    if is_blankish(cell_value):
        return [], []

    by_basename: dict[str, list[dict[str, Any]]] = {}
    by_relative: dict[str, list[dict[str, Any]]] = {}
    for entry in manifest:
        by_basename.setdefault(entry["basename"], []).append(entry)
        by_relative.setdefault(entry["relative_path"], []).append(entry)

    attachments: list[dict[str, Any]] = []
    errors: list[dict[str, Any]] = []
    for token in split_multiline(cell_value):
        token = token.replace("\\", "/").strip()
        candidates = by_relative.get(token) or by_basename.get(Path(token).name) or []
        if len(candidates) == 1:
            attachment = _entry_to_attachment(candidates[0])
            attachment["source_token"] = token
            attachments.append(attachment)
        elif not candidates:
            errors.append({"token": token, "error": "attachment_not_found"})
        else:
            errors.append(
                {
                    "token": token,
                    "error": "attachment_ambiguous",
                    "candidates": [candidate["relative_path"] for candidate in candidates],
                }
            )
    return attachments, errors
